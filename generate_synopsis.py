#!/usr/bin/env python3
"""Generate a Word document synopsis of the GenericUtility Demo Platform exercise."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x1F, 0x2E)

# -- Title Page --
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('GenericUtility Analytics Platform')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Exercise Synopsis & Technical Architecture')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Prepared: {datetime.now().strftime("%B %d, %Y")}\n').font.size = Pt(11)
meta.add_run('Platform: Azure Kubernetes Service (AKS)\n').font.size = Pt(11)
meta.add_run('Observability: Dynatrace OneAgent Full-Stack\n').font.size = Pt(11)
meta.add_run('Repository: github.com/vietpham123/utility-demo').font.size = Pt(11)

doc.add_page_break()

# -- Table of Contents placeholder --
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Objective & Scope',
    '3. Architecture Overview',
    '4. Application 1: Customer Billing Platform',
    '5. Application 2: Outage Analytics Platform',
    '6. Technology Stack',
    '7. Infrastructure & Deployment',
    '8. Observability with Dynatrace',
    '9. Interactive UI Features',
    '10. Key Metrics',
    '11. Build & Deployment Process',
    '12. Summary of Iterations',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# -- 1. Executive Summary --
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This document provides a comprehensive synopsis of the GenericUtility Analytics Platform exercise, '
    'a multi-application, polyglot microservices platform built from scratch and deployed on Azure Kubernetes '
    'Service (AKS) with Dynatrace full-stack observability. The exercise demonstrates a realistic utility '
    'company operations environment encompassing customer billing, outage management, SCADA telemetry, '
    'demand forecasting, grid topology, crew dispatch, customer notifications, and weather correlation services.'
)
doc.add_paragraph(
    'The platform comprises 2 distinct applications with a total of 21 pods running across 6 programming '
    'languages, communicating via REST APIs, gRPC, Kafka, and RabbitMQ message brokers, and backed by '
    'TimescaleDB (PostgreSQL 15), Redis 7, and persistent event-driven data pipelines.'
)

# -- 2. Objective & Scope --
doc.add_heading('2. Objective & Scope', level=1)
doc.add_paragraph(
    'The primary objective was to create a purpose-built demo environment that showcases Dynatrace\'s '
    'observability capabilities across a complex, realistic enterprise application landscape. Specific goals included:'
)
objectives = [
    'Polyglot microservice architecture spanning Node.js, .NET 6, Java 17 (Spring Boot 3), Python 3.11 (Flask), and Go 1.22',
    'Multiple communication protocols: synchronous REST, gRPC (with Protocol Buffers), and asynchronous messaging (Kafka, RabbitMQ)',
    'Real-time data simulation: SCADA sensors, meter readings, outages, crew dispatches, weather conditions, and customer notifications',
    'Interactive single-page dashboard UI with drill-down navigation, detail modals, and a geo-projected SVG map',
    'Full CI/CD pipeline: Docker multi-stage builds, Azure Container Registry, Kubernetes rolling deployments',
    'Dynatrace OneAgent instrumentation for distributed tracing, service flow mapping, log analytics, and anomaly detection',
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

# -- 3. Architecture Overview --
doc.add_heading('3. Architecture Overview', level=1)
doc.add_paragraph(
    'The platform follows a dual-application architecture deployed in separate Kubernetes namespaces, '
    'each with its own API gateway, UI, and backend microservices.'
)

doc.add_heading('High-Level Topology', level=2)
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Component'
hdr[1].text = 'Namespace'
hdr[2].text = 'Pods'
hdr[3].text = 'External IP'
for cells in [hdr]:
    for cell in cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

rows_data = [
    ('Customer Billing App', 'customer-billing', '5', '<app-endpoint>'),
    ('Outage Analytics App', 'outage-analytics', '24', '<app-endpoint>'),
    ('Total', '—', '29', '—'),
]
for rd in rows_data:
    row = table.add_row().cells
    for i, val in enumerate(rd):
        row[i].text = val

# -- 4. Customer Billing App --
doc.add_heading('4. Application 1: Customer Billing Platform', level=1)
doc.add_paragraph(
    'The Customer Billing Platform is a .NET 6 microservices application that simulates residential and '
    'commercial utility billing operations including customer account management, invoice generation, '
    'payment processing, and usage tracking.'
)

doc.add_heading('Services (5 pods)', level=2)
billing_services = [
    ('billing-ui', '.NET 6 / Nginx', 'Single-page billing dashboard with account management interface'),
    ('billing-gateway', '.NET 6 / ASP.NET', 'API gateway aggregating all billing microservices'),
    ('customer-service', '.NET 6 / ASP.NET', 'Customer account CRUD, service plans, account lookup'),
    ('invoice-service', '.NET 6 / ASP.NET', 'Invoice generation, billing cycle management, PDF generation'),
    ('payment-service', '.NET 6 / ASP.NET', 'Payment processing, payment history, refunds'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Service'
hdr[1].text = 'Technology'
hdr[2].text = 'Responsibility'
for row_data in billing_services:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

# -- 5. Outage Analytics App --
doc.add_heading('5. Application 2: Outage Analytics Platform', level=1)
doc.add_paragraph(
    'The Outage Analytics Platform is the primary demo application — a polyglot microservices system simulating '
    'a utility company\'s operational technology (OT) and information technology (IT) landscape. It includes '
    'real-time SCADA telemetry, automated outage detection, crew dispatch with RabbitMQ-driven notifications, '
    'weather correlation using a Go/gRPC service, and a rich interactive dashboard.'
)

doc.add_heading('Services (16 pods)', level=2)
analytics_services = [
    ('analytics-ui', 'Nginx / HTML+JS', 'Interactive SPA dashboard with 11 tabs, SVG map, modals, drilldowns'),
    ('analytics-gateway', 'Node.js 18', 'API gateway routing to all backend services, simulation orchestration'),
    ('outage-service', 'Node.js 18', 'Outage lifecycle management, SCADA-triggered auto-detection, resolution tracking'),
    ('usage-service', 'Node.js 18', 'Energy usage data collection, aggregation, and analytics'),
    ('scada-service', '.NET 6 / ASP.NET', 'SCADA sensor telemetry simulation (voltage, current, frequency, temperature)'),
    ('meter-data-service', 'Java 17 / Spring Boot 3', 'AMI meter data ingestion, VEE validation, anomaly detection'),
    ('grid-topology-service', 'Node.js 18', 'Hierarchical grid model: substations → feeders → transformers → service points'),
    ('reliability-service', 'Python 3.11 / Flask', 'IEEE 1366 reliability indices (SAIDI, SAIFI, CAIDI, MAIFI)'),
    ('demand-forecast-service', 'Node.js 18', 'Load forecasting, peak demand prediction, area-level capacity analysis'),
    ('crew-dispatch-service', 'Node.js 18', 'Automated crew dispatch, priority-based assignment, ETR calculation'),
    ('notification-service', 'Node.js 18', 'Multi-channel notifications (SMS, email, push, IVR) via RabbitMQ'),
    ('weather-service', 'Go 1.22 / gRPC+HTTP', 'Weather simulation, storm forecasting, outage-weather correlation'),
    ('timescaledb', 'PostgreSQL 15', 'Time-series database for all operational data'),
    ('redis', 'Redis 7', 'Caching layer and real-time state management'),
    ('kafka', 'Apache Kafka (KRaft)', 'Event streaming for SCADA telemetry and outage events'),
    ('rabbitmq', 'RabbitMQ 3.13', 'Message broker for crew dispatch → notification pipeline'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Service'
hdr[1].text = 'Technology'
hdr[2].text = 'Responsibility'
for row_data in analytics_services:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

# -- 6. Technology Stack --
doc.add_heading('6. Technology Stack', level=1)

doc.add_heading('Programming Languages (6)', level=2)
langs = [
    ('JavaScript / Node.js 18', '7 services — gateway, outage, usage, grid-topology, demand-forecast, crew-dispatch, notification'),
    ('C# / .NET 6', '4 services — scada-service + all 3 billing services'),
    ('Java 17 / Spring Boot 3', '1 service — meter-data-service'),
    ('Python 3.11 / Flask', '1 service — reliability-service'),
    ('Go 1.22', '1 service — weather-service (dual gRPC + HTTP)'),
    ('HTML/CSS/JavaScript', '2 UIs — analytics-ui and billing-ui (single-page apps)'),
]
for lang, desc in langs:
    p = doc.add_paragraph()
    run = p.add_run(f'{lang}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Communication Protocols', level=2)
protocols = [
    'REST/HTTP — Primary synchronous communication between gateway and services',
    'gRPC with Protocol Buffers — Weather service dual-protocol (gRPC port 50051 + HTTP port 8080)',
    'Apache Kafka (KRaft mode) — Asynchronous event streaming for SCADA telemetry and outage events',
    'RabbitMQ — Message queue for crew dispatch → notification service pipeline',
]
for p_text in protocols:
    doc.add_paragraph(p_text, style='List Bullet')

doc.add_heading('Data Stores', level=2)
stores = [
    'TimescaleDB (PostgreSQL 15) — Time-series optimized database with 10+ schemas for outages, SCADA, metering, grid topology, reliability, forecasts, crew dispatches, notifications, and weather data',
    'Redis 7 — In-memory caching for dashboard aggregation and real-time state',
]
for s_text in stores:
    doc.add_paragraph(s_text, style='List Bullet')

# -- 7. Infrastructure --
doc.add_heading('7. Infrastructure & Deployment', level=1)

infra_items = [
    ('Kubernetes Cluster', 'AKS or equivalent — managed Kubernetes with autoscaling node pool'),
    ('Container Registry', 'Private Docker image registry (ACR, Docker Hub, etc.)'),
    ('Build Host', 'VM or CI server for Docker builds'),
    ('Namespaces', 'outage-analytics (24 pods), customer-billing (5 pods)'),
    ('Image Pull Secret', 'Kubernetes secret for registry authentication'),
    ('Load Balancers / Ingress', 'L7 routing with TLS termination for both apps'),
    ('Dynatrace OneAgent', 'Full-stack monitoring via DaemonSet on all nodes'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Component'
hdr[1].text = 'Details'
for row_data in infra_items:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

# -- 8. Observability --
doc.add_heading('8. Observability with Dynatrace', level=1)
doc.add_paragraph(
    'The entire platform is instrumented with Dynatrace OneAgent, deployed as a DaemonSet across all AKS nodes. '
    'This provides automatic, zero-code instrumentation of all 21 pods and covers:'
)
obs_items = [
    'Distributed Tracing — End-to-end transaction visibility across all 6 languages and 4 communication protocols (REST, gRPC, Kafka, RabbitMQ)',
    'Service Flow Mapping — Automatic topology discovery showing the complete service dependency graph',
    'Log Analytics — Structured logging from all services with automatic correlation to traces and metrics',
    'Real User Monitoring — Frontend performance monitoring on both analytics and billing UIs',
    'Database Monitoring — TimescaleDB query performance, connection pool metrics, and slow query detection',
    'Container & Kubernetes Monitoring — Pod health, resource utilization, namespace-level dashboards',
    'Anomaly Detection — Davis AI-powered baseline learning and automatic problem detection',
    'Error Injection Points — Built-in configurable error rates and latency injection for chaos engineering demos',
]
for item in obs_items:
    doc.add_paragraph(item, style='List Bullet')

# -- 9. Interactive UI --
doc.add_heading('9. Interactive UI Features', level=1)
doc.add_paragraph(
    'The Analytics UI is a sophisticated single-page application with 11 navigation tabs and rich interactive features:'
)

doc.add_heading('Dashboard Tabs', level=2)
tabs = [
    ('Overview', 'Aggregated KPIs, interactive SVG territory map, summary panels for all subsystems'),
    ('SCADA Telemetry', 'Real-time sensor readings, voltage/current/frequency monitoring, alert management'),
    ('Outage Management', 'Full outage lifecycle with expandable rows, detail modals, and timeline visualization'),
    ('Meter Data (MDMS)', 'AMI meter readings, VEE validation results, anomaly detection'),
    ('Grid Topology', 'Collapsible hierarchical tree: substations → feeders → transformers → service points'),
    ('Reliability Indices', 'IEEE 1366 metrics (SAIDI, SAIFI, CAIDI, MAIFI) with 30-day history'),
    ('Demand Forecast', 'Area-level load forecasting with interactive bar charts and peak demand indicators'),
    ('Crew Dispatch', 'Crew status, active dispatches, dispatch history with expandable detail rows'),
    ('Notifications', 'Multi-channel delivery stats, event type breakdown, notification log'),
    ('Weather', 'Regional conditions, storm forecasts, weather alerts, outage-weather correlations'),
    ('Service Health', 'Real-time health check status for all backend microservices'),
]
for tab, desc in tabs:
    p = doc.add_paragraph()
    run = p.add_run(f'{tab}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Interactive Features', level=2)
features = [
    'Service Territory Map — Full eastern US SVG map with 20+ geo-projected states; service territory highlighted with blue borders; non-service states dimmed for geographic context',
    'Map Overlays — Weather radar, wind direction arrows, precipitation rings, storm mode pulsing animations, individual outage markers (color-coded by priority), crew position markers (status-colored)',
    'Clickable KPI Cards — Overview KPIs drill down to their respective sections with hover arrow indicators',
    'Panel Header Drilldowns — Click any overview panel header to navigate to the full section',
    'Breadcrumb Navigation — Drill path shown with clickable back links (e.g., Overview › Outages)',
    'Outage Detail Modal — Rich modal with full outage info, impact stats, crew assignment, and timeline visualization',
    'Crew Detail Modal — Complete crew info with active dispatch details and cross-links to outages',
    'Expandable Table Rows — Outage, crew, and SCADA tables expand inline with 3-column detail grids and action buttons',
    'Map Marker Clicks — Outage markers open detail modal, crew markers open crew modal, weather regions drill to Weather tab',
    'Cross-Entity Navigation — Dispatches link to outages, crews link to dispatches, weather correlations link to outages',
    'Keyboard Support — Escape key closes modals; click outside modal to dismiss',
    'Auto-Refresh — Dashboard data refreshes every 15 seconds with live clock display',
]
for f_text in features:
    doc.add_paragraph(f_text, style='List Bullet')

# -- 10. Key Metrics --
doc.add_heading('10. Key Metrics', level=1)
metrics = [
    ('Total Microservices', '17 (12 analytics + 5 billing)'),
    ('Infrastructure Pods', '4 (TimescaleDB, Redis, Kafka, RabbitMQ)'),
    ('Total Running Pods', '21'),
    ('Programming Languages', '6 (JavaScript, C#, Java, Python, Go, HTML/JS)'),
    ('Docker Images', '17 (all custom-built)'),
    ('Total Source Lines', '~9,400'),
    ('Kubernetes Manifests', '2 all-in-one YAML files (~1,100 lines each)'),
    ('Database Schemas', '10+ (outages, scada, metering, grid, reliability, forecast, crews, notifications, weather)'),
    ('API Endpoints', '50+ across all services'),
    ('Git Commits', '6 (iterative feature additions)'),
    ('UI Tabs', '11'),
    ('Map States Rendered', '20+ US states'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Metric'
hdr[1].text = 'Value'
for row_data in metrics:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

# -- 11. Build & Deploy --
doc.add_heading('11. Build & Deployment Process', level=1)
doc.add_paragraph('Each iteration followed a consistent build-deploy pipeline:')
steps = [
    'Code Development — Edit source files locally',
    'Transfer to Build Host — SCP or CI pipeline to build server',
    'Docker Build — Multi-stage Docker builds (e.g., .NET SDK → runtime, Maven → JRE, Go → scratch)',
    'Push to Registry — Docker push to container registry',
    'Kubernetes Deploy — kubectl rollout restart deployment to pull latest images',
    'Verification — kubectl get pods to confirm all pods Running, test endpoints',
    'Git Commit & Push — Commit changes to repository (main branch)',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

# -- 12. Summary of Iterations --
doc.add_heading('12. Summary of Iterations', level=1)
doc.add_paragraph('The platform was built incrementally across 6 major iterations:')

iterations = [
    ('Iteration 1: Foundation', 'Built both applications from scratch — Customer Billing (.NET 6, 5 services) and Outage Analytics (Node.js/Java/.NET/Python, 12 services). Established TimescaleDB, Redis, Kafka infrastructure. Deployed to AKS with Dynatrace OneAgent. Created initial dashboard UI with 8 tabs.'),
    ('Iteration 2: Crew Dispatch & Notifications', 'Added crew-dispatch-service (Node.js) with automated priority-based crew assignment, ETR calculation, and dispatch lifecycle management. Added notification-service (Node.js) with multi-channel delivery (SMS, email, push, IVR) connected via RabbitMQ. Added RabbitMQ infrastructure pod. Pod count grew to 15.'),
    ('Iteration 3: Weather Correlation Service', 'Added weather-service (Go 1.22) as the 11th analytics microservice. Implemented dual-protocol: gRPC (Protocol Buffers on port 50051) and HTTP REST (port 8080). Simulates weather conditions for 6 regions, generates storm forecasts, creates weather-outage correlations. Added Weather tab to UI. Pod count reached 16.'),
    ('Iteration 4: Service Territory Map', 'Added an interactive SVG-based service territory map to the Overview tab. Initial version with 6 simplified state polygons (IL, PA, NJ, MD, DE, DC). Weather radar overlays, wind direction arrows, precipitation rings, storm mode animations, outage count badges, and interactive tooltips.'),
    ('Iteration 5: Enhanced Map — Full US & Crew/Outage Markers', 'Expanded the map to show ~20 US states using geo-projected SVG polygons covering lat 33-45, lng -92 to -71. Service territory states highlighted with blue dashed borders; surrounding states dimmed. Added individual outage markers (priority-colored with pulsing animation for critical), crew position markers (status-colored), dispatch lines connecting crews to outages, and a comprehensive legend.'),
    ('Iteration 6: Interactive Drilldowns & Detail Modals', 'Added comprehensive interaction layer: clickable KPI cards that drill to respective tabs, panel header drilldowns with breadcrumb navigation, outage detail modal (full info + timeline + crew cross-link), crew detail modal (with active dispatch info), expandable table rows in Outages/Crew/SCADA with inline detail grids, clickable map markers (outages → modal, crews → modal, regions → weather tab), collapsible grid topology tree, and keyboard/backdrop modal dismissal.'),
]
for title, desc in iterations:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

# -- Footer --
doc.add_page_break()
doc.add_heading('Document Information', level=1)
info_items = [
    ('Author', 'Viet Pham'),
    ('Date', datetime.now().strftime('%B %d, %Y')),
    ('Platform Version', 'v1.0'),
    ('Analytics UI', '<analytics-app-endpoint>'),
    ('Billing UI', '<billing-app-endpoint>'),
    ('Kubernetes Cluster', 'Managed Kubernetes'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Field'
hdr[1].text = 'Value'
for row_data in info_items:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

# Save
output_path = '/Users/viet.pham/Library/CloudStorage/OneDrive-Dynatrace/Documents/Dynatrace Data/Exelon/GenericUtility_Platform_Synopsis.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
